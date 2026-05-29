"""
Duplique TERROIR_document_projet_final_v3.docx en v4
et remplace les sections 2 et 3 par le nouveau contenu.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
import shutil
from copy import deepcopy
from docx import Document
from docx.shared import Pt

SRC = r"C:\Users\carin\OneDrive\Downloads\TERROIR_document_projet_final_v3.docx"
DST = r"C:\Users\carin\OneDrive\Downloads\TERROIR_document_projet_final_v4.docx"

LONG_SEP  = "────────────────────────────────────────────────────────────────────────"
SHORT_SEP = "──────────────────────────────────────────────────────────────────────────"

# ── Contenu des nouvelles sections ──────────────────────────────────────────
# (texte, style, bold)
# style: "Normal" | "List Bullet"
NEW_CONTENT = [
    # ── SECTION 2 ────────────────────────────────────────────────────────────
    ("2. PROBLÈME RÉEL RÉSOLU", "Normal", True),
    ("", "Normal", False),
    ("Le constat de Phase 1 : le nord, invisible mais sous tension", "Normal", True),
    ("", "Normal", False),
    (
        "L'analyse de 23 859 événements GDELT sur l'année 2025 a révélé un paradoxe "
        "territorial précis.",
        "Normal", False,
    ),
    (
        "Le nord du Bénin — Alibori, Borgou, Atacora — représente 4 % de la couverture "
        "médiatique internationale du pays. Le sud en concentre 94,6 %. La disproportion "
        "est structurelle : les médias couvrent là où sont les institutions, les capitales, "
        "les correspondants.",
        "Normal", False,
    ),
    (
        "Mais ce que ces 4 % révèlent est significatif : le ton médiatique du nord est "
        "−4,10, contre −1,09 pour le sud. Quatre fois plus négatif. Et quand on décompose "
        "les codes CAMEO de cette couverture, la dominante est sécuritaire — tensions, "
        "coercition, incidents, conflits de basse intensité.",
        "Normal", False,
    ),
    (
        "La lecture est directe : quand le nord est couvert, c'est presque toujours parce "
        "qu'il s'est passé quelque chose de négatif sur le plan sécuritaire. Et si la "
        "couverture est faible, ce n'est pas que le territoire est calme — c'est qu'il est "
        "invisible.",
        "Normal", False,
    ),
    ("", "Normal", False),
    (SHORT_SEP, "Normal", False),
    ("", "Normal", False),
    ("Ce que l'invisibilité cache", "Normal", True),
    ("", "Normal", False),
    (
        "La sécurité interne dans le nord du Bénin ne se réduit pas aux conflits armés. "
        "Elle recouvre un spectre plus large et plus quotidien :",
        "Normal", False,
    ),
    ("vols de bétail récurrents dans les zones rurales de l'Alibori", "List Bullet", False),
    (
        "insécurité sur les axes routiers commerciaux "
        "(Parakou–Malanville, Natitingou–Tanguiéta)",
        "List Bullet", False,
    ),
    (
        "tensions communautaires liées au foncier entre agriculteurs et éleveurs, "
        "à recrudescence saisonnière",
        "List Bullet", False,
    ),
    (
        "agressions sur les agents de terrain (santé, éducation, ONG) qui se déplacent "
        "en moto vers les villages enclavés",
        "List Bullet", False,
    ),
    (
        "pressions sur les marchés hebdomadaires frontaliers",
        "List Bullet", False,
    ),
    ("", "Normal", False),
    (
        "Ces incidents ont des caractéristiques communes : ils se produisent loin des "
        "centres, ils ne remontent que rarement dans les médias nationaux ou internationaux, "
        "et ils suivent des cycles — saisonniers, économiques, politiques — que personne "
        "n'a encore eu les moyens de documenter systématiquement.",
        "Normal", False,
    ),
    (
        "Résultat : les acteurs qui opèrent dans ces zones — agents préfectoraux, "
        "coordinateurs d'ONG locales, responsables de services communaux, agents de santé "
        "communautaire — prennent leurs décisions sur la base de ce qu'ils entendent par "
        "téléphone, par bouche-à-oreille, ou par l'expérience accumulée. Pas sur la base "
        "d'un signal consolidé.",
        "Normal", False,
    ),
    ("", "Normal", False),
    (SHORT_SEP, "Normal", False),
    ("", "Normal", False),
    ("Le scénario concret", "Normal", True),
    ("", "Normal", False),
    (
        "Un coordinateur terrain d'une ONG nationale de santé communautaire basée à Kandi "
        "(Alibori) envoie chaque semaine ses agents motorisés dans des villages à 40 ou "
        "60 kilomètres. Avant chaque mission, il passe plusieurs appels pour « prendre la "
        "température » — un chef de village, un contact à la gendarmerie locale, un collègue "
        "d'une autre organisation.",
        "Normal", False,
    ),
    (
        "Cette veille informelle fonctionne. Elle repose entièrement sur un réseau humain "
        "que lui seul a construit, qu'il est le seul à pouvoir activer, et qui ne produit "
        "aucune trace exploitable par quiconque d'autre dans son organisation.",
        "Normal", False,
    ),
    (
        "Quand il part en congé, cette connaissance ne se transfère pas. Quand un incident "
        "survient sur un axe qu'il n'a pas encore eu le temps de « sonner », la décision "
        "d'envoyer ou non un agent se prend dans le flou.",
        "Normal", False,
    ),
    (
        "Ce n'est pas un manque de compétence. C'est un manque de système.",
        "Normal", False,
    ),
    ("", "Normal", False),
    (SHORT_SEP, "Normal", False),
    ("", "Normal", False),
    ("Ce que TERROIR résout", "Normal", True),
    ("", "Normal", False),
    (
        "TERROIR ne prétend pas remplacer ce coordinateur. Il lui fournit ce qu'il construit "
        "actuellement à la main, de façon automatisée et partageable : un signal consolidé "
        "sur l'état sécuritaire de son territoire, nourri à la fois par la couverture "
        "médiatique et par les remontées terrain des acteurs qui se trouvent dans les mêmes "
        "zones que lui.",
        "Normal", False,
    ),
    (
        "Quand le Score de Tension Territorial de l'Alibori monte anormalement sur quatorze "
        "jours, il reçoit une alerte. Il peut décider de reporter une mission, de contacter "
        "son réseau de façon ciblée, ou de signaler à sa direction. La décision reste "
        "humaine. Mais elle est prise sur la base d'un signal que personne n'aurait agrégé "
        "à sa place.",
        "Normal", False,
    ),
    ("", "Normal", False),

    # ── SECTION 3 ────────────────────────────────────────────────────────────
    (LONG_SEP, "Normal", False),
    ("3. UTILISATEURS ET IMPACT TERRAIN", "Normal", True),
    ("", "Normal", False),
    ("Utilisateur principal — Le coordinateur local, nord-Bénin", "Normal", True),
    ("", "Normal", False),
    (
        "Profil : responsable d'une antenne locale d'ONG nationale ou internationale, agent "
        "de service préfectoral, coordinateur de programme de développement opérant dans les "
        "départements du nord. Il n'est pas spécialiste de la sécurité. La sécurité fait "
        "partie de ses contraintes opérationnelles, pas de sa mission principale.",
        "Normal", False,
    ),
    (
        "Sa réalité : il prend des décisions qui engagent la sécurité de ses équipes "
        "plusieurs fois par semaine — envoyer ou non des agents sur un axe, confirmer ou "
        "reporter une mission en zone rurale, évaluer si une réunion communautaire peut se "
        "tenir dans un village. Il le fait sans outil dédié, avec les moyens du bord.",
        "Normal", False,
    ),
    ("", "Normal", False),
    ("Ce que TERROIR lui apporte :", "Normal", False),
    (
        "Une carte live des événements sécuritaires signalés dans sa zone — médias "
        "internationaux et remontées terrain, mis à jour toutes les quinze minutes",
        "List Bullet", False,
    ),
    (
        "Le Score de Tension Territorial de son département, calculé quotidiennement et "
        "contextualisé par rapport aux quatre-vingt-dix jours précédents",
        "List Bullet", False,
    ),
    (
        "Une alerte automatique quand le score dépasse le seuil — avec les signaux "
        "déclencheurs, le précédent historique le plus proche, et une recommandation d'action",
        "List Bullet", False,
    ),
    (
        "Un formulaire de signalement pour contribuer lui-même à la carte : ce qu'il voit "
        "sur le terrain enrichit le système pour les autres",
        "List Bullet", False,
    ),
    ("", "Normal", False),
    ("Ce que ça change concrètement :", "Normal", False),
    (
        "Il ne décide plus dans le flou. Il décide avec un signal.",
        "List Bullet", False,
    ),
    (
        "Son réseau informel reste utile — mais il devient une couche de confirmation, "
        "pas la seule source.",
        "List Bullet", False,
    ),
    (
        "Quand il s'absente, la connaissance ne part pas avec lui. Le système tient le fil.",
        "List Bullet", False,
    ),
    (
        "Les incidents qui ne remontent pas dans les médias peuvent être signalés directement "
        "— et devenir visibles pour les autres acteurs de la zone.",
        "List Bullet", False,
    ),
    ("", "Normal", False),
    (SHORT_SEP, "Normal", False),
    ("", "Normal", False),
    ("Utilisateurs secondaires", "Normal", True),
    ("", "Normal", False),
    (
        "Les agents de terrain mobiles — infirmiers itinérants, enseignants en zones rurales "
        "enclavées, agents d'inspection qui se déplacent régulièrement sur des axes non "
        "sécurisés. Ils peuvent signaler un incident depuis leur téléphone. Ils peuvent "
        "consulter la carte avant de partir.",
        "Normal", False,
    ),
    (
        "Les responsables préfectoraux — chefs de service préfectoral, sous-préfets, "
        "responsables des forces de l'ordre locales. TERROIR leur fournit une vision "
        "territoriale consolidée que les remontées administratives classiques ne produisent "
        "pas en temps réel. Le tableau de bord par département leur donne une lecture "
        "d'ensemble sans attendre les rapports hebdomadaires.",
        "Normal", False,
    ),
    (
        "Les journalistes locaux et correspondants régionaux — qui couvrent le nord avec "
        "peu de moyens. La carte live leur donne une vue des signalements récents dans leur "
        "zone de couverture. Ce qu'ils confirment sur le terrain peut alimenter le système "
        "en retour.",
        "Normal", False,
    ),
    (
        "Les citoyens et communautés locales — sans inscription obligatoire pour consulter "
        "la carte. Le formulaire de signalement est accessible à tout utilisateur. Un vol de "
        "bétail signalé par un éleveur dans un village de l'Alibori devient visible sur la "
        "carte et peut croiser un signal GDELT de même nature dans la même zone.",
        "Normal", False,
    ),
    ("", "Normal", False),
    (SHORT_SEP, "Normal", False),
    ("", "Normal", False),
    ("La chaîne d'impact", "Normal", True),
    ("", "Normal", False),
    (
        "L'impact de TERROIR se mesure à une seule chose : la qualité des décisions "
        "opérationnelles prises par des acteurs qui, aujourd'hui, décident sans signal "
        "consolidé.",
        "Normal", False,
    ),
    (
        "Pas d'évacuation précipitée parce que personne n'avait vu le signal trois jours "
        "avant. Pas de mission annulée au dernier moment parce que l'information a mis trop "
        "de temps à remonter. Pas d'incident ignoré parce qu'aucun média ne l'a couvert.",
        "Normal", False,
    ),
    (
        "Un territoire qui parle, à ceux qui ont besoin de l'entendre, au moment où ça "
        "compte encore.",
        "Normal", False,
    ),
    ("", "Normal", False),
]


def _make_run(para, text, bold=False):
    """Ajoute un run formaté à un paragraphe existant."""
    run = para.add_run(text)
    run.bold = bold
    return run


def insert_para_before(doc, anchor_elem, text, style="Normal", bold=False):
    """
    Crée un paragraphe dans le document, le déplace juste avant anchor_elem,
    puis retourne l'élément XML inséré.
    """
    if style == "List Bullet":
        try:
            p = doc.add_paragraph(style="List Bullet")
        except Exception:
            p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()

    if text:
        _make_run(p, text, bold=bold)

    new_elem = p._element
    # Déplacer avant l'ancre
    anchor_elem.addprevious(deepcopy(new_elem))
    # Supprimer l'original ajouté en fin de document
    new_elem.getparent().remove(new_elem)


def main():
    shutil.copy2(SRC, DST)
    doc = Document(DST)

    paras = doc.paragraphs
    total = len(paras)
    print(f"Paragraphes totaux : {total}")

    # Vérification des indices clés
    print(f"[23] {paras[23].text[:60]}")
    print(f"[43] {paras[43].text[:60]}")
    print(f"[44] {paras[44].text[:60]}")

    # ── 1. Mémoriser les éléments à supprimer et l'ancre ────────────────────
    to_remove = [paras[i]._element for i in range(23, 44)]
    anchor    = paras[44]._element  # séparateur avant section 4 — reste intact

    # ── 2. Supprimer sections 2 et 3 originales ─────────────────────────────
    for elem in to_remove:
        elem.getparent().remove(elem)
    print(f"Supprimé {len(to_remove)} paragraphes (sections 2 et 3 originales).")

    # ── 3. Injecter le nouveau contenu avant l'ancre ────────────────────────
    for text, style, bold in NEW_CONTENT:
        insert_para_before(doc, anchor, text, style=style, bold=bold)
    print(f"Injecté {len(NEW_CONTENT)} nouveaux paragraphes.")

    # ── 4. Sauvegarder ──────────────────────────────────────────────────────
    doc.save(DST)
    print(f"Sauvegardé : {DST}")


if __name__ == "__main__":
    main()
